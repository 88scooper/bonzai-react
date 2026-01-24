#!/usr/bin/env python3
"""
Simple script to convert markdown to Word document.
Requires: pip install python-docx markdown
"""

import sys
import os

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import markdown
    from markdown.extensions import codehilite, fenced_code
except ImportError:
    print("Error: Required packages not installed.")
    print("Please install: pip install python-docx markdown")
    sys.exit(1)

def convert_markdown_to_word(md_file, docx_file):
    """Convert markdown file to Word document."""
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Parse markdown (basic parsing)
    lines = md_content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            i += 1
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            i += 1
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            i += 1
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
            i += 1
        # Code blocks
        elif line.startswith('```'):
            # Collect code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph(''.join(code_lines))
                p.style = 'No Spacing'
                # Use monospace font for code
                for run in p.runs:
                    run.font.name = 'Courier New'
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            i += 1
        # Numbered lists
        elif line[0].isdigit() and '. ' in line[:5]:
            p = doc.add_paragraph(line.split('. ', 1)[1], style='List Number')
            i += 1
        # Regular paragraphs
        else:
            # Handle inline code and bold
            p = doc.add_paragraph()
            parts = line.split('`')
            for j, part in enumerate(parts):
                if j % 2 == 0:
                    # Regular text
                    run = p.add_run(part)
                else:
                    # Code (inline)
                    run = p.add_run(part)
                    run.font.name = 'Courier New'
            i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

if __name__ == '__main__':
    md_file = 'BONZAI_COMPREHENSIVE_SUMMARY.md'
    docx_file = 'BONZAI_COMPREHENSIVE_SUMMARY.docx'
    
    if not os.path.exists(md_file):
        print(f"Error: {md_file} not found")
        sys.exit(1)
    
    convert_markdown_to_word(md_file, docx_file)
